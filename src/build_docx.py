# -*- coding: utf-8 -*-
"""교수님 지정 과제 #2·#3 docx 생성. #3은 운영계획안 수행일지 3항목 구조(주요 논의/피드백·조치/증빙)를 따름."""
import os, sys
try: sys.stdout.reconfigure(encoding='utf-8')
except Exception: pass
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"C:\Users\WannaGoHome\Desktop\내 문서\coss\사이버보안 WE-MEET"
OUTDIR = os.path.join(ROOT, "output", "제출물")
os.makedirs(OUTDIR, exist_ok=True)
FONT = "맑은 고딕"
NAVY = RGBColor(0x18, 0x29, 0x4C)
TEAL = RGBColor(0x0C, 0x8A, 0x80)


def base_doc():
    d = Document()
    st = d.styles['Normal']; st.font.name = FONT; st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), FONT)
    return d


def H(d, text, size=15, color=NAVY, space_before=10, space_after=4):
    p = d.add_paragraph(); p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    r.font.name = FONT; r._element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), FONT)
    return p


def P(d, text, bullet=False, size=10.5, indent=0):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.25
    if indent: p.paragraph_format.left_indent = Pt(indent)
    r = p.add_run(("• " if bullet else "") + text); r.font.size = Pt(size)
    r.font.name = FONT; r._element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), FONT)
    return p


def meta_line(d, text):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x5C, 0x6B, 0x82)
    r.font.name = FONT; r._element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), FONT)


# ========== #2 진행간 어려운 점 (자유 양식) ==========
def build_q2(out=None):
    d = base_doc()
    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("사이버보안 WE-Meet — 팀 과제 #2: 진행하면서 어려웠던 점")
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    r.font.name = FONT; r._element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), FONT)
    meta_line(d, "교과목: 사이버보안 WE-Meet   |   과제: AI 기반 공격탐지 에이전트 (네트워크 침입 탐지)")
    meta_line(d, "팀: 컴퓨터융합학부 2인   |   작성일: 2026-06-26")
    P(d, "계획만 세우다가 실제로 만들고 평가해보니 생각보다 막히는 부분이 많았습니다. 저희가 부딪힌 어려움과, 교수님께 여쭙고 싶은 점을 정리했습니다.", size=10.5)

    items = [
        ("1. 정확도가 100% 나왔는데, 오히려 믿기 어려웠습니다",
         ["처음에 하루치 데이터로 모델을 만들었더니 정확도가 100%로 나왔습니다. 너무 잘 나와서 안을 들여다봤더니, 모델이 사실상 항목 하나(패킷 크기와 관련된 값 하나)만 보고 정상과 공격을 나누고 있었습니다. 공격이 실제로 어떻게 동작하는지를 배운 게 아니라, 이 데이터에서만 우연히 잘 들어맞는 값 하나에 기대고 있었던 것입니다.",
          "그래서 날짜가 다르거나 다른 종류의 공격이 들어오면 전혀 맞히지 못했습니다.",
          "여쭙고 싶은 점 — 실무에서 공개 데이터로 모델을 만들 때, 이렇게 '그 데이터에서만 잘 맞는 가짜 성능'을 어떻게 걸러내시는지, 평가를 믿어도 되는지 확인하실 때 무엇을 보시는지 궁금합니다."]),
        ("2. 처음 보는 공격은 거의 못 잡습니다",
         ["과거에 있던 공격으로 모델을 학습시키고, 그 이후에 새로 나타난 공격을 잡게 해봤더니 탐지율이 거의 0까지 떨어졌습니다. 데이터를 그냥 무작위로 섞어서 평가하면 성능이 높게 보이지만, 실제 상황처럼 '과거에 배워서 미래를 막는' 순서로 평가하면 크게 떨어졌습니다.",
          "원인을 찾아보니 모델 자체가 부족해서가 아니라, 날마다 들어오는 공격 종류가 달라서 과거에 통하던 기준이 새 공격에는 안 맞는 것이 문제였습니다. 그래서 시간 흐름 정보(같은 대상에 짧은 시간 동안 얼마나 자주·반복적으로 접속하는지 등)를 더해봤더니, 봇넷 탐지율이 거의 0에서 0.49 정도까지 올라갔습니다.",
          "여쭙고 싶은 점 — 한 번도 본 적 없는(제로데이) 공격을 모델 하나로 잡길 기대하는 게 현실적인지, 현업에서는 보통 어느 정도를 목표로 두시는지 궁금합니다. 또 방화벽·관제시스템·단말 보안처럼 여러 장비를 함께 쓰는 환경에서 AI 탐지가 실제로 맡는 비중이 어느 정도인지도 알고 싶습니다."]),
        ("3. 데이터에 빠진 정보 때문에 막히는 부분이 있습니다",
         ["저희가 쓰는 공개 데이터는 통신 상대의 IP 주소가 지워져 있습니다. 봇넷은 '같은 서버에 규칙적으로 계속 접속'하는 점이 중요한 단서인데, IP가 없으니 이 단서를 아예 만들 수가 없었습니다.",
          "웹 공격(예: SQL 삽입, 악성 스크립트 삽입)은 실제로 어떤 내용을 보냈는지가 데이터에 없고 횟수·크기 같은 숫자만 남아 있어서, 평범한 웹 접속과 구분이 잘 안 됩니다. 그래서 지금 구조로는 거의 못 잡습니다.",
          "IP와 내용이 다 들어 있는 원본 데이터는 하루치만 수십 GB, 전체로는 약 440GB라서 직접 다루기에 부담이 큽니다.",
          "여쭙고 싶은 점 — 자원이 넉넉지 않은 상황에서 IP나 요청 내용 없이도 접근할 수 있는 현실적인 방법이 있는지, 아니면 이 부분은 한계로 분명히 적어두고 넘어가는 게 맞는지 조언을 구하고 싶습니다."]),
        ("4. 위험도 점수와 알림 기준을 어떻게 정해야 할지 막막합니다",
         ["경보 하나하나에 0~100점짜리 위험 점수를 매기고, 몇 점부터 먼저 처리할지 기준선을 정해야 합니다. 저희는 시험용 데이터가 아니라 과거 데이터를 기준으로 정해야 실제와 비슷해진다고 보고, 정상을 공격으로 잘못 알리는 일(오탐)을 일정 수준 아래로 누르는 방식으로 기준을 잡았습니다.",
          "여쭙고 싶은 점 — 실제 보안관제 현장에서 분석가들이 신뢰하는 위험도·우선순위 표시 방식이 어떤 것인지, 그리고 하루에 어느 정도의 오탐까지는 감당 가능한지(현실적인 알림 양) 감을 얻고 싶습니다."]),
        ("5. 프로젝트 범위 — 교수님 피드백을 받고 방향을 정했습니다",
         ["처음에는 '이 프로젝트를 어디까지 만드는 게 적당한가'를 정하기 어려웠습니다. 네트워크 탐지 전반을 넓게 다룰지, 특정 공격에 집중할지 고민이 있었습니다.",
          "교수님께서 '탐지 강화가 목적이라면 특정 공격으로 범위를 좁혀 정확도를 높여보라'고 피드백 주셨고, 이를 받아들여 봇넷(Bot) 탐지 강화로 방향을 정했습니다. 봇넷은 처음 보는 공격으로 두면 기존 방식이 거의 못 잡는 어려운 공격이라, 정확도를 끌어올리는 의미가 큽니다.",
          "앞으로는 무작위 분할의 높은 점수에 기대지 않고, 과거→미래(날짜 교차) 평가에서 봇넷 탐지율을 높이는 것을 목표로 진행하겠습니다."]),
    ]
    for h, lines in items:
        H(d, h, size=12)
        for ln in lines:
            P(d, ln, bullet=True, indent=10)
    H(d, "마무리", size=12)
    P(d, "저희는 정확도를 높이는 것보다, 만든 결과를 '어디까지 믿어도 되는지' 정직하게 확인하는 데 초점을 두고 진행하고 있습니다. 위 다섯 가지에 대해 실무 경험에서 나온 조언을 부탁드립니다. 감사합니다.")
    out = out or os.path.join(OUTDIR, "사이버보안 WE-Meet 팀과제2 진행간 어려운점.docx")
    d.save(out); return out


# ========== #3 수행일지 1주차 (참고 1-2 양식: 3항목 구조) ==========
def build_q3():
    d = base_doc()
    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("하기 계절학기 프로젝트 수행일지 (1주차)")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
    r.font.name = FONT; r._element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), FONT)
    meta_line(d, "교과목: 사이버보안 WE-Meet   |   과제: AI 기반 공격탐지 에이전트 (네트워크 침입 탐지 및 보안로그 분석)")
    meta_line(d, "팀: 팀장(팀장, (학번 비공개)) · 팀원((학번 비공개))   |   기간: 2026-06-22 ~ 06-29   |   작성일: 2026-06-28")
    meta_line(d, "※ 6.26 온/오프라인 소집 멘토링은 운영 여건상 미진행 → 교수님 지정 과제(#1 PPT, #2 어려운점) 수행 및 비대면 피드백 요청으로 대체")

    H(d, "1. 주요 논의 주제", size=13)
    for ln in [
        "주제 구체화: ‘AI 기반 공격탐지 에이전트’를 네트워크 침입 탐지(NDR) + 보안로그 분석으로 구체화. 데이터셋 CSE-CIC-IDS2018 확정.",
        "시스템 설계: 데이터 → RF/XGBoost 이진 탐지 → 0~100 위험점수 → SHAP 근거 → LLM 자연어 설명 → Streamlit 대시보드. (탐지·점수는 ML, LLM은 설명만)",
        "정직한 평가 체계 수립: 무작위 분할의 고정확도 함정을 발견하고 날짜 교차(과거→미래) 평가로 전환.",
        "교수님께 여쭌 사항(과제 #2 연계): ①벤치마크 artifact 식별법 ②미관측 공격 일반화의 현실적 목표 ③IP/payload 부재 데이터의 대안 ④SOC 위험도·alert budget 실무 ⑤학부 프로젝트 적정 범위.",
    ]:
        P(d, ln, bullet=True, indent=10)

    H(d, "2. 검토 피드백 및 후속 조치 계획", size=13)
    P(d, "(자체 검토)", size=10.5)
    for ln in [
        "1일치 베이스라인 F1 1.000 → 단일 피처 99% 지름길 허상 확인 → 평가 설계 전면 재설계(다중 일자·누수 피처 제거·날짜 교차).",
        "미관측 공격 탐지율 붕괴 → 원인이 ‘전이(분포 변화)’임을 진단 → 시간맥락 피처 추가로 봇넷 0.009→0.49 등 개선.",
    ]:
        P(d, ln, bullet=True, indent=10)
    P(d, "(후속 조치 — 차주)", size=10.5)
    for ln in [
        "교수님 비대면 피드백(과제 #2 답변) 반영 → 평가·범위 방향 조정",
        "위험점수·임계값 운영 검증 마무리 및 대시보드 시연 자료화",
        "결과보고서 본문 정리(양식 채움 완료, 교수님 피드백 반영 예정)",
        "(여건 시) 원본 데이터 재처리로 호스트 기반 피처 추가 검토",
    ]:
        P(d, ln, bullet=True, indent=10)

    H(d, "3. 증빙 자료", size=13)
    for ln in [
        "팀 과제 #1: 프로젝트 소개용 PPT (사이버보안 WE-Meet 팀과제1 PPT.pptx, 9매)",
        "팀 과제 #2: 진행간 어려운 점 정리 (사이버보안 WE-Meet 팀과제2 진행간 어려운점.docx)",
        "코드·실험 산출물: 탐지·평가·위험점수·SHAP·LLM 설명·대시보드·로그 분석 (지표 JSON, SHAP 플롯, 대시보드 화면)",
        "※ 6.26 소집 멘토링 미진행으로 대면 회의 사진 증빙은 없으며, 비대면 메일 송부 내역으로 갈음.",
    ]:
        P(d, ln, bullet=True, indent=10)

    out = os.path.join(OUTDIR, "사이버보안 WE-Meet 팀과제3 수행일지 1주차.docx")
    d.save(out); return out


if __name__ == "__main__":
    print("저장:", build_q2())
    print("저장:", build_q3())
